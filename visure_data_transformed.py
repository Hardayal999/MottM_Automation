import re
import sys
from copy import deepcopy
from datetime import datetime
from pathlib import Path
from typing import Iterator, List, Optional, Union

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Requirement id alone: TEST_0020 / UID_1964
REQ_ID_ONLY_RE = re.compile(r'^((?:TEST|UID)_\d+)\s*$', re.I)

# Requirement id + inline text on same line: TEST_00010 some text
REQ_ID_INLINE_RE = re.compile(r'^((?:TEST|UID)_\d+)(?:\s+(.*?))?\s*$', re.I | re.S)

# Word heading styles
HEADING_NUM_RE = re.compile(r'^Heading\s+(\d+)$', re.I)

Block = Union[Paragraph, Table]
REQ_LEVELS_SUPPORTED = (2, 3, 4, 5)  # Heading 2..5 => Req Level 1..4


def iter_block_items(parent) -> Iterator[Block]:
    """Yield paragraphs and tables in document order."""
    if isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        parent_elm = parent.element.body if hasattr(parent.element, 'body') else parent._element

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def clear_body_keep_sect(doc: Document) -> None:
    """Clear document body but keep section properties."""
    body = doc._body._element
    sectPr = body.sectPr
    for child in list(body):
        body.remove(child)
    if sectPr is not None:
        body.append(sectPr)


def append_paragraph_copy(dst_doc: Document, src_para: Paragraph) -> None:
    dst_doc._body._element.insert(-1, deepcopy(src_para._p))


def append_table_copy(dst_doc: Document, src_tbl: Table) -> None:
    dst_doc._body._element.insert(-1, deepcopy(src_tbl._tbl))


def heading_level(para: Paragraph) -> Optional[int]:
    style_name = para.style.name if para.style is not None else ''
    m = HEADING_NUM_RE.match(style_name)
    return int(m.group(1)) if m else None


def normalize_text(text: str) -> str:
    return text.replace('\xa0', ' ').strip()


def set_cell_border_none(cell: _Cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        elem = tcBorders.find(qn(f'w:{edge}'))
        if elem is None:
            elem = OxmlElement(f'w:{edge}')
            tcBorders.append(elem)
        elem.set(qn('w:val'), 'nil')


def set_table_borders_none(table: Table) -> None:
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        elem = tblBorders.find(qn(f'w:{edge}'))
        if elem is None:
            elem = OxmlElement(f'w:{edge}')
            tblBorders.append(elem)
        elem.set(qn('w:val'), 'nil')

    for row in table.rows:
        for cell in row.cells:
            set_cell_border_none(cell)


def remove_cell_default_paragraph(cell: _Cell) -> None:
    if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text and len(cell.tables) == 0:
        p = cell.paragraphs[0]
        p._element.getparent().remove(p._element)


def copy_nested_table(dst_cell: _Cell, src_tbl: Table, font_size_pt: int = 11) -> None:
    """Copy a source table into a destination cell as a nested borderless table."""
    nt = dst_cell.add_table(rows=len(src_tbl.rows), cols=len(src_tbl.columns))
    nt.alignment = WD_TABLE_ALIGNMENT.LEFT

    try:
        nt.autofit = False
    except Exception:
        pass

    for i, row in enumerate(src_tbl.rows):
        for j, src_cell in enumerate(row.cells):
            dst = nt.cell(i, j)
            dst.text = src_cell.text
            for p in dst.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for r in p.runs:
                    r.font.size = Pt(font_size_pt)

    set_table_borders_none(nt)


def add_num_instance(doc: Document, abstract_num_id: int = 1) -> int:
    """
    Create a fresh numbering instance so Req Level 1 restarts at (a)
    for each Heading 1 section.
    """
    numbering = doc.part.numbering_part.element
    nums = numbering.findall(qn('w:num'))
    max_num = max((int(n.get(qn('w:numId'))) for n in nums), default=0)
    new_id = max_num + 1

    num = OxmlElement('w:num')
    num.set(qn('w:numId'), str(new_id))

    abs_id = OxmlElement('w:abstractNumId')
    abs_id.set(qn('w:val'), str(abstract_num_id))
    num.append(abs_id)

    for ilvl in range(4):
        lvl_override = OxmlElement('w:lvlOverride')
        lvl_override.set(qn('w:ilvl'), str(ilvl))

        start_override = OxmlElement('w:startOverride')
        start_override.set(qn('w:val'), '1')

        lvl_override.append(start_override)
        num.append(lvl_override)

    numbering.append(num)
    return new_id


def set_paragraph_numbering(p: Paragraph, num_id: int, ilvl: int) -> None:
    """Force paragraph numbering to use a specific numbering instance and level."""
    pPr = p._p.get_or_add_pPr()
    existing = pPr.find(qn('w:numPr'))
    if existing is not None:
        pPr.remove(existing)

    numPr = OxmlElement('w:numPr')

    ilvl_el = OxmlElement('w:ilvl')
    ilvl_el.set(qn('w:val'), str(ilvl))

    numId_el = OxmlElement('w:numId')
    numId_el.set(qn('w:val'), str(num_id))

    numPr.append(ilvl_el)
    numPr.append(numId_el)
    pPr.append(numPr)


def split_into_sections(blocks: List[Block]) -> List[List[Block]]:
    """
    Split document into sections using Heading 1 as the divider.
    Anything before first Heading 1 stays as-is.
    """
    sections: List[List[Block]] = []
    current: List[Block] = []
    started = False

    for block in blocks:
        if isinstance(block, Paragraph) and block.style.name == 'Heading 1':
            if current:
                sections.append(current)
            current = [block]
            started = True
        else:
            if started:
                current.append(block)
            else:
                sections.append([block])

    if current:
        sections.append(current)

    return sections


def parse_requirement_start(block: Paragraph):
    """
    Return parsed requirement start info or None.
    Supports:
    TEST_0020
    TEST_00010 Some inline text
    UID_1964
    Only Heading 2..5 are treated as requirement levels.
    """
    lvl = heading_level(block)
    if lvl not in REQ_LEVELS_SUPPORTED:
        return None

    txt = normalize_text(block.text)
    m = REQ_ID_INLINE_RE.match(txt)
    if not m:
        return None

    req_id = m.group(1)
    inline_text = (m.group(2) or '').strip()

    return {
        'id': req_id,
        'level': lvl - 1,   # Heading 2..5 => Req Level 1..4
        'inline_text': inline_text,
        'blocks': [],
    }


def section_has_requirements(sec: List[Block]) -> bool:
    for block in sec[1:]:
        if isinstance(block, Paragraph) and parse_requirement_start(block):
            return True
    return False


def apply_style_or_fallback(p: Paragraph, desired_style: str, style_names: set, fallback: str = 'Normal') -> None:
    if desired_style in style_names:
        p.style = desired_style
    elif fallback in style_names:
        p.style = fallback


def add_requirement_table(
    dst_doc: Document,
    reqs: list,
    section_num_id: int,
    left_width_cm: float = 14.2,
    right_width_cm: float = 3.0,
    font_size_pt: int = 11
) -> None:
    """
    Add a borderless 2-column requirement table:
      left = requirement text with Req Level style
      right = requirement ID
    """
    table = dst_doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    try:
        table.autofit = False
    except Exception:
        pass

    style_names = {s.name for s in dst_doc.styles}

    for req in reqs:
        row = table.add_row()
        left, right = row.cells
        left.width = Cm(left_width_cm)
        right.width = Cm(right_width_cm)

        remove_cell_default_paragraph(left)
        remove_cell_default_paragraph(right)

        first_block_done = False

        # If requirement had inline text on same line as TEST_xxxx
        if req.get('inline_text'):
            p = left.add_paragraph(req['inline_text'])
            apply_style_or_fallback(p, f"Req Level {req['level']}", style_names)
            set_paragraph_numbering(p, section_num_id, req['level'] - 1)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for r in p.runs:
                r.font.size = Pt(font_size_pt)
            first_block_done = True

        # Add continuation paragraphs/tables under same requirement
        for block in req['blocks']:
            if isinstance(block, Paragraph):
                txt = normalize_text(block.text)
                if not txt:
                    continue

                p = left.add_paragraph(txt)

                if not first_block_done:
                    # First content paragraph gets Req Level style and numbering
                    apply_style_or_fallback(p, f"Req Level {req['level']}", style_names)
                    set_paragraph_numbering(p, section_num_id, req['level'] - 1)
                    first_block_done = True
                else:
                    # Later continuation paragraphs keep original style if available
                    original_style = block.style.name if block.style is not None else 'Normal'
                    apply_style_or_fallback(p, original_style, style_names)

                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for r in p.runs:
                    r.font.size = Pt(font_size_pt)

            elif isinstance(block, Table):
                copy_nested_table(left, block, font_size_pt=font_size_pt)

        # Safety fallback if there was no content after the ID
        if not first_block_done:
            p = left.add_paragraph('')
            apply_style_or_fallback(p, f"Req Level {req['level']}", style_names)
            set_paragraph_numbering(p, section_num_id, req['level'] - 1)

        # Right-hand ID cell
        p_id = right.add_paragraph(req['id'])
        if 'Normal' in style_names:
            p_id.style = 'Normal'
        p_id.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for r in p_id.runs:
            r.font.size = Pt(font_size_pt)

    set_table_borders_none(table)


def transform(
    src_path: str,
    out_path: str,
    left_width_cm: float = 14.2,
    right_width_cm: float = 3.0,
    font_size_pt: int = 11
) -> None:
    """
    Main transformation:
      - preserve document shell/styles/headers/sections
      - keep non-requirement text as-is
      - convert requirement blocks into borderless 2-column tables
      - apply Req Level 1..4 styles
      - restart (a) under each Heading 1 section
    """
    src = Document(src_path)
    out = Document(src_path)  # preserve styles, numbering defs, headers, footers, etc.
    clear_body_keep_sect(out)

    blocks = list(iter_block_items(src))
    sections = split_into_sections(blocks)

    for sec in sections:
        first = sec[0]

        # Anything before first Heading 1, or non-standard sections: copy as-is
        if not (isinstance(first, Paragraph) and first.style.name == 'Heading 1'):
            for block in sec:
                if isinstance(block, Paragraph):
                    append_paragraph_copy(out, block)
                else:
                    append_table_copy(out, block)
            continue

        # Keep Heading 1 section heading
        append_paragraph_copy(out, first)

        # If section has no requirements, copy everything as-is
        if not section_has_requirements(sec):
            for block in sec[1:]:
                if isinstance(block, Paragraph):
                    append_paragraph_copy(out, block)
                else:
                    append_table_copy(out, block)
            continue

        # New numbering instance so Req Level 1 restarts at (a) per Heading 1 section
        section_num_id = add_num_instance(out, abstract_num_id=1)

        reqs = []
        current_req = None

        for block in sec[1:]:
            if isinstance(block, Paragraph):
                parsed = parse_requirement_start(block)

                if parsed:
                    if current_req:
                        reqs.append(current_req)
                    current_req = parsed
                    continue

                if current_req:
                    current_req['blocks'].append(block)
                else:
                    append_paragraph_copy(out, block)

            else:
                if current_req:
                    current_req['blocks'].append(block)
                else:
                    append_table_copy(out, block)

        if current_req:
            reqs.append(current_req)

        if reqs:
            add_requirement_table(
                out,
                reqs,
                section_num_id,
                left_width_cm=left_width_cm,
                right_width_cm=right_width_cm,
                font_size_pt=font_size_pt,
            )

    out.save(out_path)


def generate_output_path(input_path: str) -> str:
    p = Path(input_path)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return str(p.with_name(f"{p.stem}_output_{timestamp}.docx"))


def main() -> None:
    if len(sys.argv) < 2:
        print('Usage: python visure_transform_final_test.py "C:\\path\\input.docx"')
        sys.exit(1)

    input_file = sys.argv[1]
    input_path = Path(input_file)

    if not input_path.exists():
        print(f"File not found: {input_file}")
        sys.exit(1)

    if input_path.suffix.lower() != ".docx":
        print("Input file must be a .docx file")
        sys.exit(1)

    output_file = generate_output_path(input_file)

    print(f"Input : {input_file}")
    print(f"Output: {output_file}")

    transform(
        src_path=input_file,
        out_path=output_file,
        left_width_cm=14.2,
        right_width_cm=3.0,
        font_size_pt=11,
    )

    print("Done.")


if __name__ == "__main__":
    main()